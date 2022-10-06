#!/usr/bin/python3
from cgi import print_directory
from codeop import CommandCompiler
from datetime import date
from distutils.log import info
import os
import shutil
from xml.dom.minidom import Document
import cv2
import termios
from tkinter import *
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from functools import partial
from tkinter import font
from tkinter import RIGHT, TOP, ttk
from tkinter.messagebox import askokcancel, showinfo, WARNING
from tkinterdnd2 import *

import sqlite3
import functools
import glob
from datetime import *

from turtle import color, left, width
from ttkwidgets.autocomplete import AutocompleteCombobox
from PIL import ImageTk,Image
from ftplib import FTP
from os import listdir


# - Efface le 'cache' photos -
removingfiles = glob.glob('*.jpg')
for i in removingfiles:
    os.remove(i)


# Classe de 'stockage' :
class Data():

	def __init__(self) :
		self.utility="Stockage"


# Classe pour le bouton d'acceuil :
class ToggledFrame(tk.Frame):

    def __init__(self, parent, text="", *args, **options):
        tk.Frame.__init__(self, parent, *args, **options)

        self.show = tk.IntVar()
        self.show.set(0)

        self.title_frame = ttk.Frame(self)
        self.title_frame.pack(fill="x", expand=1)

        ttk.Label(self.title_frame, text=text).pack(side="left", fill="x", expand=1)

        self.toggle_button = ttk.Checkbutton(self.title_frame, width=2, command=self.toggle,variable=self.show, style='Toolbutton')
        self.toggle_button.pack(side="left")

        self.sub_frame = tk.Frame(self, relief="sunken", borderwidth=1)

    def toggle(self):
        if bool(self.show.get()):
            self.sub_frame.pack(fill="x", expand=1)
            self.toggle_button.configure(text='-')
        else:
            self.sub_frame.forget()
            self.toggle_button.configure(text='+')

# - Variable de stockage de chemins -
chemin=""
chemin_drag=""
esttuple=False

# Classe de mise en forme :
class main():


	def __init__(self):
		self.fontFil = open("sett.txt","r+")
		size = self.fontFil.read()
		self.fontFil.close()


		# - Script de création de la BDD -
		def creer_base():
			sql =["CREATE TABLE Clients (idClient INTEGER PRIMARY KEY NOT NULL UNIQUE, firmeNom TEXT, firmeAdresse TEXT, firmeCP TEXT, firmeVille TEXT, clientNom TEXT, clientPrenom TEXT, adresseMail TEXT, telephone TEXT, soldeCompte REAL, numSiret INTEGER, numTVA INTEGER);","CREATE TABLE Commandes (idCommande INTEGER PRIMARY KEY NOT NULL UNIQUE, dateDepot TEXT, adresseLivraison TEXT, cp INTEGER, ville TEXT, prixConvenu REAL,devis REAL,acompte REAL,resterPaye REAL, description TEXT, travaf TEXT, photo TEXT, matiere TEXT, poid INTEGER, idClient INTEGER NOT NULL, FOREIGN KEY (idClient) REFERENCES Clients(idClient));"]
			li = sqlite3.connect("base.db")

			if li:
				cons = li.cursor()

				for i in range(len(sql)):
					cons.execute(sql[i])
				li.commit()
				li.close()


		# - Permet de créer la BDD si elle n'exite pas déjà -
		def verifier_base(base):
			from os.path import isfile,isdir
			from os import makedirs
			if not isfile(base) and not isdir('factures'):
				creer_base()
				makedirs('factures')
			else:
				pass



		# - Appel de la fonction verifier_base à l'initialisation de l'UI -
		verifier_base('base.db')

		def toggleFullScreen(event):
			self.fullScreenState = not self.fullScreenState
			window.attributes("-fullscreen", self.fullScreenState)

		def quitFullScreen(event):
			self.fullScreenState = False
			window.attributes("-fullscreen", self.fullScreenState)
			window.geometry('1500x1000')


		barr_rech = Data()
		barr_rechC = Data()
		barr_rechCli = Data()

		window = TkinterDnD.Tk()
		window.attributes('-fullscreen', True)
		fullScreenState = False
		window.bind('<Escape>', quitFullScreen)
		window.bind('<Control-p>', toggleFullScreen)
		window.title("Logiciel Gestion Commande")
		window.geometry('2000x1000')
		window.configure(bg="#666363")

		infow = window.winfo_screenwidth()
		print("infow :"+str(infow))

		self.titrerestor = tk.Label(window,text="Gestion",font =('normal',20,'bold'),bg="#666363",).pack(fill=X)

		# - Affiche tout les clients -
		def showAllClients():
			self.topShowAllClients = tk.Toplevel()
			self.topShowAllClients.geometry("1920x1080")
			self.topShowAllClients.bind("<Escape>", lambda even: self.topShowAllClients.destroy())

			arbre = ttk.Treeview(self.topShowAllClients, column=("c1","c2", "c3","c4","c5","c6","c7","c8","c9"), show='headings')

			conn = sqlite3.connect('base.db')
			cur = conn.cursor()
			rqt = "SELECT firmeNom, firmeAdresse, firmeCP, firmeVille, clientNom, clientPrenom, adresseMail, telephone, soldeCompte FROM Clients;"
			cur.execute(rqt)
			rows = cur.fetchall()
			for row in rows:
				arbre.insert("", tk.END, values=row)
				arbre.bind("<Double-1>",showClients)
				arbre.bind("<Button-2>", unset_item)
				arbre.bind("<Button-3>", unset_item)
			cur.close()
			conn.close()

			l=["firmeNom","firmeAdresse","firmeCP","firmeVille","clientNom","clientPrenom","adresseMail","telephone","soldeCompte"]

			for i in range(len(l)):
				arbre.column("#"+str(i+1), anchor=tk.CENTER)
				arbre.heading("#"+str(i+1), text=l[i])
			arbre.pack(side = "bottom", expand=True, fill='y')

			barr_rechCli.queryString = AutocompleteCombobox(self.topShowAllClients,width=30,font=('Times', 18),completevalues=recuperer("client"))
			barr_rechCli.queryString.pack(padx=10,pady=130)
			barr_rechCli.queryString.bind("<Return>", lambda even, brci = barr_rechCli, wC = self.topShowAllClients: search(brci,wC,"client"))

			self.topShowAllClients.bind("<Return>", lambda even, brci = barr_rechCli, wC = self.topShowAllClients: search(brci,wC,"client"))

		# - Permet d'effacer un client ET UNE COMMANDE !-
		def unset_item(event):
			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem,'values')
			if (len(val) == 7 and val[0].isdigit()!=True) or len(val) == 9:
				answer = askokcancel(
				title='Confirmation',
				message='Êtes vous sûr de vouloir supprimer ce client ?',
				icon=WARNING
				)
				if answer:
					nom = val[4]
					prenom = val[5]
					conn = sqlite3.connect('base.db')
					cur = conn.cursor()
					cur.execute('DELETE FROM Commandes WHERE idClient = (SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?);', (nom, prenom) )
					conn.commit()
					cur.execute('DELETE FROM Clients WHERE clientNom = ?;', [nom])
					conn.commit()
					conn.close()
					showinfo(
					title='Suppression',
					message='Le client à été supprimé.'
					)
			else:
				answer = askokcancel(
				title='Confirmation',
				message='Êtes vous sûr de vouloir supprimer cette commande ?',
				icon=WARNING
				)
				if answer:
					idComm = val[0]
					conn = sqlite3.connect('base.db')
					cur = conn.cursor()
					cur.execute('DELETE FROM Commandes WHERE idCommande = ?', [idComm])
					conn.commit()
					conn.close()
					showinfo(
					title='Suppression',
					message='La commande à été supprimée.'
					)

		# - Recherche un client par son nom de famille ou une commande par son numéro -
		def search(entree, window, CI):
			from datetime import datetime
			import re
			conn = sqlite3.connect("base.db")
			try :
				if datetime.strptime(entree.queryString.get(), "%d-%m-%Y") is not None:
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.dateDepot = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-2>", unset_item)
						tree.bind("<Button-3>", unset_item)
					cur.close()
					return 1

			except:
				pass

			try:
				reg = re.compile(".*\..*")
				if reg.match(str(entree.queryString.get())):
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.prixConvenu = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)
					cur.close()

					return 1
			except:
				pass

			try:
				reg = re.compile(".* .*")
				if reg.match(str(entree.queryString.get())):
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE Commandes.description = ? or Commandes.travaf = ?;", (entree.queryString.get(),entree.queryString.get()))
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)
					cur.close()


					return 1
			except:
				pass


			try:
				if type(int(entree.queryString.get()))==int and CI == "general" or CI == "commande":
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7"), show='headings')
					l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient WHERE idCommande = ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-3>", unset_item)
					cur.close()

			except ValueError:
				if CI == "general" or CI == "client":
					tree = ttk.Treeview(window, column=("c1", "c2", "c3","c4","c5","c6","c7","c8","c9"), show='headings')
					l = ["firmeNom","firmeAdresse","firmeCP","firmeVille","clientNom","clientPrenom","adresseMail","telephone","soldeCompte"]
					for i in range(len(l)):
						tree.column("#"+str(i+1), anchor=tk.CENTER)
						tree.heading("#"+str(i+1), text=l[i])
					tree.place(x=0,y=288,height=500)
					tree.delete(*tree.get_children())
					cur = conn.cursor()
					cur.execute("SELECT firmeNom, firmeAdresse, firmeCP, firmeVille, clientNom, clientPrenom, adresseMail, telephone, soldeCompte FROM Clients WHERE clientNom LIKE ?;", [entree.queryString.get()])
					fetch = cur.fetchall()
					for res in fetch:
						tree.insert('', 'end', values=(res))
						tree.bind("<Double-1>", nouvCommande)
						tree.bind("<Button-3>", unset_item)
						tree.bind("<Button-2>", unset_item)
					cur.close()
				conn.close()

		def nouvCommande(event):
			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem)
			Comm = []
			if len(val['values']) == 9:
				nom = val["values"][4]
				prenom = val["values"][5]
				conn = sqlite3.connect('base.db')
				cur = conn.cursor()
				cur.execute('SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?;', (nom,prenom))
				conn.commit()
				idCli = cur.fetchall()
				print("Idc (1) : "+str(idCli))
				cur.execute('SELECT firmeAdresse,firmeCP,firmeVille FROM Clients WHERE idClient = ?;', [idCli[0][0]])
				Comm = cur.fetchall()
				windowCommande(self, idCli, Comm)

			elif len(val['values']) == 8 or len(val['values']) == 7:
				conn = sqlite3.connect('base.db')
				cur = conn.cursor()
				cur.execute('SELECT * FROM Commandes WHERE idCommande = ?;', [val['values'][0]])
				conn.commit()
				fetch = cur.fetchall()
				val = fetch[0]
				idCLI = fetch[0][12]
				windowCommande(self, idCLI, val)


		# - Récupération des noms des clients -
		def recuperer(CI):
			conn = sqlite3.connect("base.db")
			cur = conn.cursor()
			cur.execute('SELECT clientNom FROM Clients')
			fetch = cur.fetchall()
			cur.close()
			l = [r[0] for r in fetch]

			cur = conn.cursor()
			cur.execute('SELECT idCommande FROM Commandes')
			fetch1 = cur.fetchall()
			cur.close()
			conn.close()
			l1 = [str(r[0]) for r in fetch1]

			if CI == "general":
				return l+l1
			elif CI == "client":
				return l
			elif CI == "commande":
				return l1

		# - Barre de recherche et tableau -
		barr_rech.queryString = AutocompleteCombobox(window,width=30,font=('Times', 18),completevalues=recuperer("general"))
		barr_rech.queryString.pack(padx=10,pady=130)

		barr_rech.queryString.bind("<Return>", lambda even, brci = barr_rech, w = window: search(brci,w,"general"))
		window.bind("<Return>", lambda even, brci = barr_rech, w = window: search(brci,w,"general"))

		# - Affiche tout les clients -
		def showAllCommandes():

			self.topShowAllCommandes = tk.Toplevel()
			self.topShowAllCommandes.geometry("1920x1080")
			self.topShowAllCommandes.bind("<Escape>", lambda even: self.topShowAllCommandes.destroy())

			arbre = ttk.Treeview(self.topShowAllCommandes, column=("c1", "c2","c3","c4","c5","c6","c7","c8"), show='headings')
			conn = sqlite3.connect('base.db')
			cur = conn.cursor()
			cur.execute("SELECT idCommande,Clients.clientNom,dateDepot,cp,ville,devis,description,travaf FROM Commandes JOIN Clients ON Commandes.idClient=Clients.idClient ;")
			rows = cur.fetchall()
			for row in rows:
				arbre.insert("", tk.END, values=row)
				arbre.bind("<Double-1>", nouvCommande)
				arbre.bind("<Button-3>", unset_item)
				arbre.bind("<Button-2>", unset_item)
			cur.close()
			conn.close()


			l = ["Numéro de la commande","Nom du client","Date de dépôt","Code postale","Ville","Devis","Description","Travail à faire"]

			for i in range(len(l)):
				arbre.column("#"+str(i+1), anchor=tk.CENTER)
				arbre.heading("#"+str(i+1), text=l[i])
			arbre.pack(side="bottom",expand=True, fill='y')


			barr_rechC.queryString = AutocompleteCombobox(self.topShowAllCommandes,width=30,font=('Times', 18),completevalues=recuperer("commande"))
			barr_rechC.queryString.pack(padx=10,pady=130)

			barr_rechC.queryString.bind("<Return>", lambda even, brci = barr_rechC, wC = self.topShowAllCommandes: search(brci,wC,"commande"))
			#self.topShowAllCommandes.bind("<Return>", lambda even, brci = barr_rechC, wC = self.topShowAllCommandes: search(brci,wC,"commande"))



		# - Permet d'ajouter une commande -
		def windowCommande(self, idC, Comm):
			from datetime import datetime
			color = "#52575B"
			data = Data() # Instance de Data() qui va nous permettre de stocker les str récupérées grâce à main()
			self.top = tk.Toplevel()
			self.top.title("Nouvelle Commande")
			self.top.geometry("800x700")
			self.top.configure(background=color)
			self.top.bind("<Escape>", lambda d: self.top.destroy())
			def afficImage(event):
				testvariable.set(event.data)
    			# get the value from string variable
				self.top.file_name=testvariable.get()
				global chemin_drag
				chemin_drag=self.top.file_name
				# takes path using dragged file
				image_path = Image.open(str(self.top.file_name))
				# resize image
				reside_image = image_path.resize((300, 205), Image.ANTIALIAS)
				# displays an image
				self.top.image = ImageTk.PhotoImage(reside_image)
				image_label = Label(labelframe, image=self.top.image).pack()
			def prendrephoto():
				i=0
				cap = cv2.VideoCapture(0)
				while(True):
					ret, frame = cap.read()
					rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2BGRA)

					cv2.imshow('frame', rgb)
					if cv2.waitKey(1) & 0xFF == ord('a'):
						if os.path.isfile("/Applications/restor/"+str(i)+'.jpg')==True:
							while os.path.isfile("/Applications/restor/"+str(i)+'.jpg')==True:
								i+=1
							out = cv2.imwrite(str(i)+'.jpg', frame)
							break
						else:
							out = cv2.imwrite(str(i)+'.jpg', frame)
							break
				global chemin
				chemin=str(i)+".jpg"
				cap.release()
				cv2.destroyAllWindows()
			def enregistrerLocal(chemin_photo,i):
				file=chemin_photo
				ftp = FTP('localhost')
				ftp.login(user='john',passwd='john')
				ftp.cwd("../../srv/ftp") # On se déplace vers le dossier où l'on veux stocker nos photos
				if os.path.isfile("/srv/ftp/"+str(i)+'.jpg')==True:
					while os.path.isfile("/srv/ftp/"+str(i)+'.jpg')==True:
						i+=1
					ftp.storbinary("STOR "+ str(i)+".jpg", open(file, "rb"))
				else:
					ftp.storbinary("STOR "+ str(i)+".jpg", open(file, "rb"))
				global chemin
				chemin=str(i)+".jpg"
				ftp.quit()
			def enregistrer(chemin_photo,idC):
				chemin_photo.save(str(idC)+".jpg")
				windowCommande()

			def setCommtoBon(data,idC,Comm):
				con = sqlite3.connect('base.db')
				cur = con.cursor()
				cur.execute('SELECT * FROM Clients WHERE idClient = ?;',[Comm[12]])
				infoc = cur.fetchall()

				'''info Clients'''
				nom_ent = infoc[0][1]
				adress = infoc[0][2]
				cp=infoc[0][3]
				ville=infoc[0][4]
				prenom=infoc[0][5]
				nom=infoc[0][6]
				mail=infoc[0][7]
				tel = infoc[0][8]
				sir=infoc[0][10]
				tva=infoc[0][11]

				'''info Commande'''
				num_commande=Comm[0]
				date=Comm[1]
				prix = Comm[5]
				devis=Comm[6]
				desc=Comm[9]
				travail_a_faire=Comm[10]

				qnt = str(1)
				tva=int(devis)*0.2
				totale = 'a faire'
				ttc = 'a faire'

				'''tableau 1 & 2'''

				sett1 = ("[total]","[ttc]","[tva]","[travail_a_f]","[description]","[date]","[datefdm]","[num]","[Entrecli]","[Adresse de l'entreprise du client]","[SIRET du client]","[Numéro de TVA du client]","[tel_cli]","[cp]","[ville]","[qnt]","[prix]","[devis]","[Nom de l'entreprise]")
				sett2 = (totale,ttc,tva,travail_a_faire,desc,date,datetime.today(),num_commande,nom_ent,adress,sir,tva,tel,cp,ville,qnt,prix,devis,nom_ent,nom_ent)

				'''Copy et creation du fichier de bon de commande'''

				src = 'exempl.docx'
				dest = '/Users/mathismottet/Desktop'+Comm[0]+'.docx'
				shutil.copy2(src,dest)

				'''Boucle d'affectation de variable'''
				i = 0
				while sett1[i] != '\0':
					file = open("exemple1.docx", "r",encoding="utf-8")
					replacement = ""
					for line in file:
						line = line.strip()
						changes = line.replace(sett1[i],str(sett2[i]))
						replacement = replacement + changes + "\n"
					file.close()
					# opening the file in write mode
					fout = open("/Users/mathismottet/Desktop"+Comm[0]+".docx", "w")
					fout.write(replacement)
					fout.close()
					i = i + 1
			stad = 0
			global esttuple
			if type(Comm) == tuple:
				stad = 1
				self.top.geometry("1500x1070")
				esttuple=True
				self.titredateDepot = tk.Label(self.top,text="Date de dépôt",bg=color,font =('normal',20,'bold')).grid(row=1)
				self.titreidCommande = tk.Label(self.top,text="Identifiant de la commande",bg=color,font =('normal',20,'bold')).grid(row=2)
				self.titreAdresseLivraison = tk.Label(self.top,text="Adresse de livraison",bg=color,font =('normal',20,'bold')).grid(row=3)
				self.titreCP = tk.Label(self.top,text="Code postale",bg=color,font =('normal',20,'bold')).grid(row=4)
				self.titreVille = tk.Label(self.top,text="Ville",bg=color,font =('normal',20,'bold')).grid(row=5)
				self.titrePrixConvenue = tk.Label(self.top,text="Prix Convenu",bg=color,font =('normal',20,'bold')).grid(row=6)
				self.titreDevis= tk.Label(self.top,text="Devis",bg=color,font =('normal',20,'bold')).grid(row=7)
				self.titreAcompte = tk.Label(self.top,text="Acompte",bg=color,font =('normal',20,'bold')).grid(row=8)
				self.titreDescription = tk.Label(self.top, text="Description",bg=color,font =('normal',20,'bold')).grid(row=9)
				self.titretravaf = tk.Label(self.top, text="Travail à faire", bg=color,font =('normal',20,'bold')).grid(row=10)

				data.dateDepot = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[1]!=None:
					data.dateDepot.insert(0, Comm[1])
					data.dateDepot.config(state='disabled')
				data.idCommande = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[0]!=None:
					data.idCommande.insert(0, Comm[0])
					data.idCommande.config(state='disabled')
				data.adresseLivraison = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[2]!=None:
					data.adresseLivraison.insert(0, Comm[2])
					data.adresseLivraison.config(state='disabled')
				data.cp = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[3]!=None:
					data.cp.insert(0, Comm[3])
					data.cp.config(state='disabled')
				data.ville = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[4]!=None:
					data.ville.insert(0, Comm[4])
					data.ville.config(state='disabled')
				data.prixConvenu = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[5]!=None:
					data.prixConvenu.insert(0, Comm[5])
				data.devis = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[6]!=None:
					data.devis.insert(0, Comm[6])
				data.acompte = tk.Entry(self.top,font =('normal',20,'bold'))
				if Comm[7]!=None:
					data.acompte.insert(0, Comm[7])
				data.Description = tk.Text(self.top,height=8,width=40,font =('normal',20,'bold'))
				if Comm[9]!=None:
					data.Description.insert(END, Comm[9])
				data.travaf = tk.Text(self.top,height=8,width=40,font =('normal',20,'bold'))
				if Comm[10]!=None:
					data.travaf.insert(END, Comm[10])

				#data.Compte = tk.Entry(self.top,width=40,font =('normal',20,'bold'))
				#if Comm[10]!=None:
				#	data.Compte.insert(END, Comm[10)]
				data.dateDepot.grid(row=1, column=2)
				data.idCommande.grid(row=2, column=2)
				data.adresseLivraison.grid(row=3, column=2)
				data.cp.grid(row=4, column=2)
				data.ville.grid(row=5, column=2)
				data.prixConvenu.grid(row=6, column=2)
				data.devis.grid(row=7, column=2)
				data.acompte.grid(row=8, column=2)
				#data.Compte.grid(row =9, column=2)
				data.Description.grid(row=9, column=2)
				data.travaf.grid(row=10, column=2)
                
			else:
				esttuple=False
				stad = 2
				self.titreidCommande = tk.Label(self.top,text="Identifiant de la commande",bg=color,font =('normal',20,'bold')).grid(row=1)
				self.titreDescription = tk.Label(self.top, text="Description",bg=color,font =('normal',20,'bold')).grid(row=2)
				self.titretravaf = tk.Label(self.top, text="Travail à faire", bg=color,font =('normal',20,'bold')).grid(row=3)

				data.dateDepot = tk.Entry(self.top,width=30,font =('normal',20,'bold'),highlightthickness=2)
				data.dateDepot.insert(0,str(date.today()))
				data.idCommande = tk.Entry(self.top,width=30,font =('normal',20,'bold'),highlightthickness=2)
				data.Description = tk.Text(self.top,height=5,width=30,font =('normal',20,'bold'))
				data.travaf = tk.Text(self.top,height=5,width=30,font =('normal',20,'bold'))

				data.dateDepot.grid(row=1,column=2)
				data.idCommande.grid(row=1, column=2)
				data.Description.grid(row=2, column=2)
				data.travaf.grid(row=3, column=2)

			# - Permet l'enregistrement de la BDD -
			def step1toGetVar(stad,data,self,idC,Comm):
				if stad == 1:
					getVarCommande(data,self,idC,stad,Comm)
					setCommtoBon(data,idC,Comm)
					self.top.destroy
				else:
					getVarCommande(data,self,idC,stad,Comm)

			def getVarCommande(data,self,idC,stad,Comm):
				global chemin
				if esttuple==True and chemin_drag!="":
					enregistrerLocal(chemin_drag,idC)

				if esttuple==True and chemin!="" and chemin_drag=="":
					enregistrerLocal(chemin,idC)

				photo="/srv/ftp/"+chemin
				if photo=="/srv/ftp/":
					photo=""

				existe=False
				if stad == 2 :
					if data.idCommande.get()!="" and data.Description.get("1.0",END)!="" and data.travaf.get("1.0",END)!="":
						existe=False
						conn = sqlite3.connect('base.db')
						cur = conn.cursor()
						cur.execute('SELECT idCommande FROM Commandes')
						conn.commit()
						fetch = cur.fetchall()
						idComm = data.idCommande.get()
						for i in fetch:
							a = i[0] - int(idComm)
							if a == 0:
								existe=True
								data.idCommande.configure(highlightbackground = "red", highlightcolor= "red",fg="red")
					
					if existe==False:
						rqt = "INSERT OR REPLACE INTO Commandes (dateDepot,idClient,idCommande,description,travaf,adresseLivraison,cp,ville) VALUES ('"+data.dateDepot.get()+"','"+str(idC[0][0])+"','"+data.idCommande.get()+"','"+data.Description.get("1.0",END)+"','"+data.travaf.get("1.0",END)+"','"+str(Comm[0][0])+"','"+str(Comm[0][1])+"','"+str(Comm[0][2])+"')"
						cur.execute(rqt)
						conn.commit()
						cur.close()
						conn.close()
						self.top.destroy()
						windowCommande(self, idC, Comm)


				else :
					if data.idCommande.get()!="" and data.Description.get("1.0",END)!="" and data.travaf.get("1.0",END)!="":
						conn = sqlite3.connect('base.db')
						cur = conn.cursor()
						cur.execute('SELECT idCommande FROM Commandes')
						conn.commit()
						fetch = cur.fetchall()


						if existe==False:
							if esttuple==True:
								cur.execute("INSERT OR REPLACE INTO Commandes (dateDepot,idCommande,adresseLivraison,cp,ville,prixConvenu,devis,acompte,resterPaye,description,travaf,idClient) VALUES ('"+data.dateDepot.get()+"','"+str(int(data.idCommande.get()))+"','"+data.adresseLivraison.get()+"','"+data.cp.get()+"','"+data.ville.get()+"','"+data.prixConvenu.get()+"','"+data.devis.get()+"','"+data.acompte.get()+"','"+str((float(data.devis.get()))-(float(data.acompte.get())))+"','"+data.Description.get("1.0",END)+"','"+data.travaf.get("1.0",END)+"','"+str(idC)+"')")
								conn.commit()
								cur.close()
								conn.close()
								self.top.destroy()
					else :
						'''Mettre les contour en rouge si pas remplie'''
						pass


			self.top.bind("<Return>", lambda even,d=data,s=self,i=idC: step1toGetVar(stad,data,self,idC,Comm))
			self.btnWebcam = tk.Button(self.top, text='Prendre une photo', highlightbackground=color,command=prendrephoto).place(x=695,y=250)

		def fact_cli():
			from docx import Document
			from datetime import datetime
			con = sqlite3.connect('base.db')

			cur = con.cursor()
			cur.execute('SELECT idClient FROM Commandes;')
			idCli = cur.fetchall()

			l =[]

			for i in range(len(idCli)):
				for e in range(len(idCli[i])):
					l.append(str(idCli[i][e]))
			l = list(set(l))

			comm = []
			for i in range(len(l)):
				cur.execute('SELECT * FROM Commandes WHERE idClient=?;', l[i])
				comm.append(cur.fetchall())

			for i in range(len(comm)):
				doc = Document()
				for e in range(len(comm[i])):
					for x in range(len(comm[i][e])):
						l.append(str(comm[i][e][x]))
					doc.add_paragraph(' '.join(l))
					l = []
				doc.save('factures/' + str(comm[i][0][11]) + '_' +datetime.today().strftime("%d-%m-%Y")+'.docx')
			showinfo(
			title='Facture',
			message='Les factures ont été générées.'
			)

		def setCommtoBon(data):
			from datetime import datetime
			from docx import Document
			from os import listdir

			con = sqlite3.connect('base.db')

			cur = con.cursor()
			cur.execute('SELECT * FROM Clients WHERE idClient = ?;', [data.idClient.get()])

			infoc = cur.fetchall()

			nom_entcli = infoc[0][1]
			adr_cli = data.adresseLivraison.get()
			adr_ent = infoc[0][2]
			tel_cli = infoc[0][8]
			prix = data.prixConvenu.get()
			qnt = 1

			entre = {"[date]":datetime.today().strftime('%d-%m-%Y'),"[datefdm]":"30-" + datetime.today().strftime('%m-%Y'),"[Nom de l'entreprise]":"REST'OR","[Entrecli]":nom_entcli,"[Adresse de l'entreprise du client]":adr_cli, "[SIRET du client]":"SRTcli", "[Numéro de TVA du client]":"VTA000", "[Adresse]":"ROUTE DE PONT L'EVEQUE 27260 CORMEILLES", "[SIRET]":"SRT000", "[Numéro de TVA]":"TV090909", "[Téléphone]":tel_cli, "[QNT]":qnt, "[prix]":prix}

			doc = Document('factures/exemple.docx')

			for i in entre:
				for p in doc.paragraphs:
					if p.text.find(i)>=0:
						p.text = p.text.replace(i, entre[i])

			dirs = listdir("factures/")
			dirs.sort(reverse=True)

			for i in range(len(dirs)):
				if "Facture" in dirs[i]:
						indice = dirs[i][len(dirs[i]) -6]
						break
				else:
						indice = 0
			doc.save('factures/Facture'+ str(int(indice) + 1)+'.docx')
			showinfo(
			title='Bon de commande',
			message='Le bon de commande à été généré.'
			)



		# - Permet d'ajouter un client -
		def windowClient(self):
			color = "#52575B"
			data = Data()

			self.top = tk.Toplevel(bg=color)
			self.top.title("Nouveau client")
			self.top.geometry("400x400")
			self.top.bind("<Escape>", lambda d: self.top.destroy())
			#self.labelTitre = tk.Label(self.top,text="Nouveaux clients",fg= color,bg=color,font=10).grid(row=0)
			self.titrefirmeNom = tk.Label(self.top,text='Nom Entreprise',font =('normal',22,'bold'),bg=color).grid(row=2)
			self.titrefirmeAdresse = tk.Label(self.top,text='Adresse Entreprise',font =('normal',22,'bold'),bg=color).grid(row=3)
			self.titrefirmeCP = tk.Label(self.top,text='Code Postale Entreprise',font =('normal',22,'bold'),bg=color).grid(row=4)
			self.titrefirmeVille = tk.Label(self.top,text='Ville Entreprise',font =('normal',22,'bold'),bg=color).grid(row=5)

			self.titreclientNom = tk.Label(self.top,text='Nom',font =('normal',22,'bold'),bg=color).grid(row=6)
			self.titrePrenom = tk.Label(self.top,text='Prenom',font =('normal',22,'bold'),bg=color).grid(row=7)
			self.titreadresseMail = tk.Label(self.top,text='Adresse Mail',font =('normal',22,'bold'),bg=color).grid(row=8)
			self.titreSoldeCompte = tk.Label(self.top,text='Compte poids',font =('normal',22,'bold'),bg=color).grid(row=9)
			self.titreTelephone = tk.Label(self.top,text='Tel',font =('normal',22,'bold'),bg=color).grid(row=10)
			self.titreNumSiret = tk.Label(self.top,text='Numéro de siret',font =('normal',22,'bold'),bg=color).grid(row=11)
			self.titreNumTVA = tk.Label(self.top,text='Numéro de TVA',font =('normal',22,'bold'),bg=color).grid(row=12)

			data.firmeNom = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.firmeAdresse = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.firmeCP = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.firmeVille = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")

			data.clientNom = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.clientPrenom = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.adresseMail = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.soldeCompte = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.telephone = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.numSiret = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")
			data.numTVA = tk.Entry(self.top,font =('normal',22,'bold'),width=25,highlightthickness=2,highlightbackground = color,relief="solid")

			data.firmeNom.grid(row=2, column=2)
			data.firmeAdresse.grid(row=3, column=2)
			data.firmeCP.grid(row=4, column=2)
			data.firmeVille.grid(row=5, column=2)
			data.clientNom.grid(row=6, column=2)
			data.clientPrenom.grid(row=7, column=2)
			data.adresseMail.grid(row=8, column=2)
			data.soldeCompte.grid(row=9, column=2)
			data.telephone.grid(row=10, column=2)
			data.numSiret.grid(row=11, column=2)
			data.numTVA.grid(row=12, column=2)


			# - Permet l'enregistrement de la BDD -
			def getVarClients(data,self):
				if data.firmeNom.get() == "":
					data.firmeNom.config(highlightbackground = "red", highlightcolor= "red")
				if data.clientNom.get() == "":
					data.clientNom.config(highlightbackground = "red", highlightcolor= "red")

				if data.clientNom.get() != "":
					data.clientNom.config(highlightbackground = color, highlightcolor= color)
				if data.firmeNom.get() != "":
					data.firmeNom.config(highlightbackground =color, highlightcolor= color)

				if data.firmeNom.get()!="" and data.clientNom.get()!="":
					data.firmeNom.config(highlightbackground = color, highlightcolor= color)
					data.clientNom.config(highlightbackground = color, highlightcolor= color)
					conn = sqlite3.connect('base.db')
					cur = conn.cursor()
					req = "SELECT idClient FROM Clients WHERE idClient = (SELECT MAX(idClient) FROM Clients)"
					cur.execute(req)
					idCo = cur.fetchall()
					a = idCo[0][0]
					a =+ 1
					rqt = "INSERT INTO Clients (idClient,firmeNom,firmeAdresse,firmeCP,firmeVille,clientNom,clientPrenom,adresseMail,telephone,soldeCompte, numSiret, numTVA) VALUES ('"+str(a)+"','"+data.firmeNom.get()+"','"+data.firmeAdresse.get()+"','"+data.firmeCP.get()+"','"+data.firmeVille.get()+"','"+data.clientNom.get()+"','"+data.clientPrenom.get()+"','"+data.adresseMail.get()+"','"+data.telephone.get()+"','"+data.soldeCompte.get()+"','"+data.numSiret.get()+"','"+data.numTVA.get()+"');"
					cur = conn.cursor()
					cur.execute(rqt)
					conn.commit()
					cur.close()
					conn.close()
					self.top.destroy()

			self.top.bind("<Return>", lambda e, d=data,s=self: getVarClients(data,self))
			self.btnValide = tk.Button(self.top, text='Valider', highlightbackground =color,command=partial(getVarClients,data,self)).place(x=650,y= 430)


		def showClients(event):

			tree = event.widget
			citem = tree.focus()
			val = tree.item(citem)
			color1 = "#666363"
			fg = 13
			if len(val['values']) == 9:
				entreprise = val["values"][0]
				nom = val["values"][4]
				prenom = val["values"][5]
				compte = val["values"][3]
				tel = val["values"][7]
				mail = val["values"][6]

				self.topShowCommande = tk.Toplevel(bg=color1)
				self.topShowCommande.geometry("3000x700")
				self.topShowCommande.bind("<Escape>", lambda d: self.topShowCommande.destroy())

				conn = sqlite3.connect('base.db')
				cur = conn.cursor()

				tree = ttk.Treeview(self.topShowCommande,column=("c1", "c2", "c3","c4","c5","c6","c7","c8","c9","c10","c11"), show='headings')

			# - IL AFFICHE TOUTES LES COLONNES
				l=["Identifiant de la commande","Date de dépôt","Adresse de livraison","Code Postale","Ville","Prix convenu", "Devis", "Acompte", "Reste à payer", "Description", "Identifiant du client"]

				for i in range(len(l)):
					tree.column("#"+str(i+1), anchor=tk.CENTER)
					tree.heading("#"+str(i+1), text=l[i])
				tree.place(x=0,y=250)

				cur.execute('SELECT * FROM Commandes WHERE idClient = (SELECT idClient FROM Clients WHERE clientNom = ? and clientPrenom = ?);', (nom,prenom) )
				fetch = cur.fetchall()
				for res in fetch:
					tree.insert('', 'end', values=(res))
					tree.bind("<Button-2>", unset_item) #test
					tree.bind("<Button-3>", unset_item) #test
				cur.close()
				conn.close()

				entreCompte = 100
				compte = str(entreCompte) + 'g'

				self.LabelFrames = tk.Label(self.topShowCommande,text=entreprise,font = fg,bg=color1).place(x=20,y=25)
				self.LabelName = tk.Label(self.topShowCommande,text=nom,bg=color1).place(x=20,y=50)
				self.LabelFirstName = tk.Label(self.topShowCommande,text=prenom,bg=color1).place(x=20,y=75)

				self.LabelMail = tk.Label(self.topShowCommande,text=mail,bg=color1).place(x=20,y=110)
				self.LabelNumTel = tk.Label(self.topShowCommande,text=tel,bg=color1).place(x=20,y=135)

				self.LabelCompte = tk.Label(self.topShowCommande,text=compte,bg=color1).place(x=340,y=135)
				self.titreCompte = tk.Label(self.topShowCommande,text='Compte poids : ',bg=color1).place(x=220,y=135)

				btnModif = Button(self.topShowCommande,text="modifier",width=3,height=1,font=20,highlightbackground = color1,bg = 'grey').place(x=700)

		def windowSett():
			self.settFrame = Frame(window,bg="white",width=400,height=10000,borderwidth = 1,relief="raised",).pack(side=RIGHT)

			self.btnFlech = tk.Button(self.settFrame,text="➡️",width=3,height=1,font=20,highlightbackground = "white",bg = 'grey').place(x=0,y=0)

		def multiBtn(w):
			def destroyMultiBtn(w):
				btn1.config(command=partial(multiBtn,w),text='+')
				btnprolong2.destroy()
				btnprolong3.destroy()
				btnprolong4.destroy()
			btn1.config(command=partial(destroyMultiBtn,w),text='-')
			btnprolong2 = tk.Button(window,text="Nouveau client",bg="#666363",relief="ridge",highlightbackground = "#666363",command = partial(windowClient,self),width=18)
			btnprolong2.place(x=w,y=105)
			btnprolong3 = tk.Button(window,text="Toutes les commandes",bg="#666363",relief="ridge",highlightbackground = "#666363",command = showAllCommandes,width=18)
			btnprolong3.place(x=w,y=130)
			btnprolong4 = tk.Button(window,text="Tout les clients",bg="#666363",relief="ridge",highlightbackground = "#666363",command = showAllClients,width=18)
			btnprolong4.place(x=w,y=155)
		infow = infow - 300
		btn2 = tk.Button(window,text="Fermer",width=5,height=1,font=20,highlightbackground = "#666363",command=quit).place(x=infow,y=30)
		btn1 = Button(window, relief="raised",text="+",borderwidth=1,width=3,height=1,highlightbackground = "#666363",command=partial(multiBtn,infow))
		btn1.place(x=infow,y=80)
		# - Affiche la fenêtre principale à l'écran puis attend que l'usager pose une action -
		window.mainloop()


# - Initialisation de l'UI -
app = main()
